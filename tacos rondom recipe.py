import docx
from  PIL import Image, ImageDraw, ImageFont  # this imports the image and image font when i'm doing my project for the image part
import requests
from docx import Document
from docx.shared import Inches
#  Brian Hullan ITEC 1150-60. This is my final project code tha will generate a recipe book of taco recipes including a rondom pictur of a tco and three rondomly genrated recipes.



# my first step is to open the image that i'll use for my poject. to do that i'm using image = image.open to open it.
image = Image.open('christine-siracusa-vzX2rgUbQXM-unsplash.jpg')    # this opens the image that will use for the project.

# the next step i'm going to resize my picture not more than 800 or less
image.thumbnail((800, 800))                  # resizing image to 800 using image.thubnall this will resize it and use   (()) to put the numbers between


# the next step is going to be image .show()  so that it shows my image when i run it



# next step i'm going  to draw a text around my image

draw = ImageDraw.Draw(image)

# for next one i'm going to use font color for text on the image. so i downloaded DjaVusans from the website

font = ImageFont.truetype('DejaVuSans (1).ttf', 40)    # this will create the font text on the image and choose 40 for the font size


# 40 as for font size

# for the draw text around my image i'm using draw.text to draw text on my image
draw.text([10,475], 'Random Taco Cookbook', fill= 'red', font = font)   # for that is going to be Random Taco Cookbook
# use 10, 475 to draw the text and the text will show on the image.
# for the color i choose red.


image.show()  # this shows the image when every thing is typed and resized the picture it will show the outcome

#  since i finished everything i'm going to go head and save my picture so that it appears on the project side

image.save('Rondom tacos recipe .jpeg')   # Saving the image as this will show up on the project side and i will be able to access it

# this next seciton of code will create the data for the Rondom Taco cookbook and the image on word document

# for my next step i'm going to create lists to store dictionaries in

Three_tacos =[]    # this list i will be using to creat dictionaries for my three taco recipes

# my other step would be going over the loop to put and take the Three tacos for my project

for i in range(3):    # i'm useing a for loop here to get the result for the url three tacos recipe

        Three_tacos_recipes_url_data = requests.get('https://taco-1150.herokuapp.com/random/?full_taco=true').json()   #
        Three_tacos .append(Three_tacos_recipes_url_data) # this one is added list for the three tacos
        print(Three_tacos)  # this will print the three tacos

# this next section  of the library let';s me create word document
word_document = docx.Document()  # start with a new blank page

word_document.add_paragraph('Rondom Taco Cookbook', 'Title')  # add a title for Rondom taco recipe


word_document.add_picture('Rondom tacos recipe .jpeg',width=Inches(5.0), height=Inches(6.78))  # adding the image on word


word_document.add_paragraph('Christine Siracusa')  # this will be printed under the picture

word_document.add_paragraph('https://unsplash.com/photos/vzX2rgUbQXM')  # this is the url where i get the picture from

word_document.add_paragraph('Mohamed Adan')  # the code was created by me

word_document.add_page_break()   # this is going to be a pager break for the all five components of the recipe.


# for my next step is going ta have all five components of the recipe. and loop over them

for i in range(len(Three_tacos)):    # this looping is going to pull the five taco recipes
        word_document.add_paragraph(Three_tacos[0]["seasoning"]["recipe"])

        # the first three tacos going to print will be seasoning recipe
        word_document.add_paragraph(Three_tacos[0]["condiment"]["recipe"])

        # the second one is going to be condiment recipe
        word_document.add_paragraph(Three_tacos[0]["mixin"]["recipe"])

        # the third one is going to be the mixin recipe
        word_document.add_paragraph(Three_tacos[0]["base_layer"]["recipe"])

        # the fourth one is going to be base layer
        word_document.add_paragraph(Three_tacos[0]["shell"]["recipe"])

        # the last one is going to be shell recipe



        #  save my word document


word_document.save('all five components of the recipe..docx')   # this is going to save the word document








